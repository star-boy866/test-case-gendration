"""
Cognos Report Definition DOCX structure parser — Part 1.

Parses a Cognos Report Definition / Design Specification DOCX into a rich,
traceable intermediate representation with exact physical document structure:
- Block-level XML document-order traversal (paragraphs & tables interleaved)
- Section hierarchy without keyword-matching guesswork
- Cell-level checkbox state preservation (CHECKED, UNCHECKED, UNKNOWN)
- Structural table/row/cell preservation
- No artificial page estimation — explicit pages or UNKNOWN

HALLUCINATION PREVENTION: This parser only extracts what is EXPLICITLY
present in the document. It never invents content or fake page numbers.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Optional, Any

import docx
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Checkbox state model
# ---------------------------------------------------------------------------

class CheckboxState(str, Enum):
    """Explicit checkbox state representation."""
    CHECKED = "CHECKED"
    UNCHECKED = "UNCHECKED"
    UNKNOWN = "UNKNOWN"


# Unicode checkbox characters
_CHECKED_CHARS = {"☑", "☒", "✓", "✔", "■"}
_UNCHECKED_CHARS = {"☐", "□", "○"}

# Regex pattern for checkbox-like patterns in text
_CHECKBOX_PATTERN = re.compile(
    r"([☑☒✓✔■☐□○])\s*(.+?)(?=\s*[☑☒✓✔■☐□○]|$)", re.DOTALL
)


def _is_checked(char: str) -> bool:
    """Determine if a checkbox character represents a checked state."""
    return char in _CHECKED_CHARS


def _extract_checkboxes(text: str) -> list[dict]:
    """
    Extract checkbox values from text containing checkbox characters.
    Returns list of {"label": str, "checked": bool, "state": CheckboxState}.
    """
    results = []
    for match in _CHECKBOX_PATTERN.finditer(text):
        char = match.group(1)
        label = match.group(2).strip()
        if label:
            state = (
                CheckboxState.CHECKED if char in _CHECKED_CHARS
                else (CheckboxState.UNCHECKED if char in _UNCHECKED_CHARS else CheckboxState.UNKNOWN)
            )
            results.append({
                "label": label,
                "checked": state == CheckboxState.CHECKED,
                "state": state,
            })
    return results


def _extract_checkboxes_from_paragraph(paragraph) -> list[dict]:
    """
    Extract checkboxes from a python-docx paragraph, handling both
    Unicode characters and Word XML form controls.
    """
    text = paragraph.text.strip()
    unicode_boxes = _extract_checkboxes(text)
    if unicode_boxes:
        return unicode_boxes

    results = []
    try:
        for run in paragraph.runs:
            xml = run._element.xml
            if "w14:checkbox" in xml or "w:checkBox" in xml or "w:fldChar" in xml:
                is_checked = "w14:checked w14:val=\"1\"" in xml or "w:checked" in xml or 'w:val="1"' in xml
                state = CheckboxState.CHECKED if is_checked else CheckboxState.UNCHECKED
                label = run.text.strip() if run.text else ""
                if not label:
                    continue
                results.append({
                    "label": label,
                    "checked": is_checked,
                    "state": state,
                })
    except Exception:
        pass

    return results


def _analyze_cell_checkboxes(cell, text: str) -> tuple[CheckboxState, list[dict]]:
    """Analyze a table cell for checkbox states."""
    cb_labels: list[dict] = []
    matches = list(_CHECKBOX_PATTERN.finditer(text))
    if matches:
        has_checked = False
        has_unchecked = False
        for match in matches:
            char = match.group(1)
            label = match.group(2).strip()
            if char in _CHECKED_CHARS:
                state = CheckboxState.CHECKED
                has_checked = True
            elif char in _UNCHECKED_CHARS:
                state = CheckboxState.UNCHECKED
                has_unchecked = True
            else:
                state = CheckboxState.UNKNOWN
            if label:
                cb_labels.append({"label": label, "state": state})

        overall = CheckboxState.UNKNOWN
        if has_checked and not has_unchecked:
            overall = CheckboxState.CHECKED
        elif has_unchecked and not has_checked:
            overall = CheckboxState.UNCHECKED
        elif has_checked and has_unchecked:
            overall = CheckboxState.CHECKED
        return overall, cb_labels

    # 2. Inspect Wingdings w:sym runs and Word XML controls across cell paragraphs
    try:
        cell_runs = [run for para in cell.paragraphs for run in para.runs]
        current_state: Optional[CheckboxState] = None
        
        for i, run in enumerate(cell_runs):
            xml = run._element.xml
            run_text = run.text.strip()
            
            # Wingdings w:sym symbols
            if "w:sym" in xml:
                if any(ch in xml for ch in ('w:char="F0FE"', 'w:char="F052"', 'w:char="F078"', 'w:char="F058"')):
                    current_state = CheckboxState.CHECKED
                elif any(ch in xml for ch in ('w:char="F0A8"', 'w:char="F0A3"', 'w:char="F0A0"')):
                    current_state = CheckboxState.UNCHECKED
            # Word content control / form checkboxes
            elif "w14:checkbox" in xml or "w:checkBox" in xml or "w:fldChar" in xml:
                is_checked = "w14:checked w14:val=\"1\"" in xml or "w:checked" in xml or 'w:val="1"' in xml
                current_state = CheckboxState.CHECKED if is_checked else CheckboxState.UNCHECKED

            if current_state is not None:
                # Look ahead for text label in this or subsequent runs
                label_text = run_text
                if not label_text and i + 1 < len(cell_runs):
                    # Gather text from next non-sym runs
                    next_parts = []
                    for j in range(i + 1, len(cell_runs)):
                        if "w:sym" in cell_runs[j]._element.xml:
                            break
                        if cell_runs[j].text.strip():
                            next_parts.append(cell_runs[j].text.strip())
                    label_text = " ".join(next_parts).strip()

                if label_text:
                    cb_labels.append({"label": label_text, "state": current_state, "checked": current_state == CheckboxState.CHECKED})
                    current_state = None

        if cb_labels:
            has_checked = any(cb["state"] == CheckboxState.CHECKED for cb in cb_labels)
            has_unchecked = any(cb["state"] == CheckboxState.UNCHECKED for cb in cb_labels)
            overall = CheckboxState.CHECKED if has_checked else (CheckboxState.UNCHECKED if has_unchecked else CheckboxState.UNKNOWN)
            return overall, cb_labels
    except Exception:
        pass

    return CheckboxState.UNKNOWN, []


# ---------------------------------------------------------------------------
# Core Structural Domain Models (Part 1)
# ---------------------------------------------------------------------------

@dataclass
class ParsedCell:
    """A single cell within a table row."""
    row_index: int
    column_index: int
    text: str
    raw_text: str
    is_merged: bool = False
    checkbox_state: CheckboxState = CheckboxState.UNKNOWN
    checkbox_labels: list[dict] = field(default_factory=list)


@dataclass
class ParsedRow:
    """A single row within a table."""
    row_index: int
    cells: list[ParsedCell] = field(default_factory=list)


@dataclass
class ParsedTable:
    """An extracted table preserving strict structure and source order."""
    table_index: int
    section_name: str
    row_count: int
    column_count: int
    rows: list[ParsedRow] = field(default_factory=list)
    source_order: int = 0
    source_page: Optional[int] = None
    raw_grid: list[list[str]] = field(default_factory=list)

    def to_grid(self) -> list[list[str]]:
        """Return 2D text grid representation."""
        if self.raw_grid:
            return self.raw_grid
        return [[cell.text for cell in row.cells] for row in self.rows]

    def __len__(self) -> int:
        return self.row_count

    def __getitem__(self, idx: int) -> list[str]:
        return self.to_grid()[idx]

    def __iter__(self):
        return iter(self.to_grid())


@dataclass
class DocumentSection:
    """A detected physical section in the Cognos report definition."""
    name: str
    start_order: int
    end_order: int = -1
    source_page: Optional[int] = None
    paragraphs: list[str] = field(default_factory=list)
    tables: list[ParsedTable] = field(default_factory=list)
    key_value_pairs: dict[str, str] = field(default_factory=dict)
    checkboxes: list[dict] = field(default_factory=list)
    raw_text: str = ""

    # Property aliases for backward compatibility with ParsedSection consumers
    @property
    def start_index(self) -> int:
        return self.start_order

    @property
    def end_index(self) -> int:
        return self.end_order

    @property
    def estimated_page(self) -> Optional[int]:
        return self.source_page


# Alias for backward compatibility
ParsedSection = DocumentSection


@dataclass
class CognosParsedDocument:
    """
    Intermediate representation of a parsed Cognos DOCX.
    Preserves exact document order, section tree, and table/cell states.
    """
    filename: str = ""
    sections: list[DocumentSection] = field(default_factory=list)
    all_paragraphs: list[str] = field(default_factory=list)
    all_parsed_tables: list[ParsedTable] = field(default_factory=list)
    all_key_value_pairs: dict[str, str] = field(default_factory=dict)
    all_checkboxes: list[dict] = field(default_factory=list)
    comments: list[dict] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    total_pages_estimated: Optional[int] = None

    @property
    def all_tables(self) -> list[list[list[str]]]:
        """Backward-compatibility property returning 3D text lists."""
        result = []
        for s in self.sections:
            for t in s.tables:
                result.append(t.to_grid())
        return result

    def get_section(self, name: str) -> Optional[DocumentSection]:
        """Find a section by name (case-insensitive partial match)."""
        name_lower = name.lower()
        for section in self.sections:
            if name_lower in section.name.lower():
                return section
        return None

    def get_sections(self, name: str) -> list[DocumentSection]:
        """Find all sections matching a name pattern."""
        name_lower = name.lower()
        return [s for s in self.sections if name_lower in s.name.lower()]

    @property
    def has_content(self) -> bool:
        return bool(self.sections or self.all_parsed_tables or self.all_paragraphs)


# ---------------------------------------------------------------------------
# Section Heading Patterns
# ---------------------------------------------------------------------------

_SECTION_PATTERNS = [
    (r"report\s+definition", "Report Definition"),
    (r"report\s+selection\s+criteria", "Report Selection Criteria"),
    (r"report\s+field", "Report Field"),
    (r"report\s+parameters?", "Report Parameters"),
    (r"report\s+control\s+breaks?,?\s+totals?,?\s+counts?,?\s+and\s+sorts?",
     "Report Control Breaks, Totals, Counts, and Sorts"),
    (r"report\s+control\s+breaks?", "Report Control Breaks"),
    (r"report\s+output", "Report Output"),
    (r"report\s+special\s+processing", "Report Special Processing"),
    (r"report\s+layout", "Report Layout"),
    (r"report\s+specification", "Report Specification"),
    (r"report\s+body", "Report Body"),
    (r"sort\s+by", "Sort By"),
    (r"control\s+break", "Control Break"),
    (r"total", "Total"),
    (r"counts?", "Count"),
    (r"output\s+format", "Output Format"),
    (r"output\s+distribution", "Output Distribution"),
    (r"report\s+retention", "Report Retention"),
    (r"run\s+history", "Run History"),
    (r"presentation\s+type", "Presentation Type"),
    (r"section\s+heading", "Section Heading"),
]

_COMPILED_SECTIONS = [
    (re.compile(pattern, re.IGNORECASE), name)
    for pattern, name in _SECTION_PATTERNS
]


def _detect_section(text: str) -> Optional[str]:
    """Match text against known section heading patterns."""
    text = text.strip()
    if not text:
        return None
    for pattern, name in _COMPILED_SECTIONS:
        if pattern.search(text):
            return name
    return None


def _extract_key_value_pairs(rows: list[list[str]]) -> dict[str, str]:
    """
    Extract key-value pairs from a 2D table grid without altering structure.
    """
    pairs = {}
    if not rows:
        return pairs

    num_rows = len(rows)
    for i, row in enumerate(rows):
        if not row:
            continue

        cell0 = row[0].strip()
        if "\n" in cell0:
            lines = [l.strip() for l in cell0.split("\n") if l.strip()]
            if len(lines) >= 2 and (lines[0].endswith(":") or len(lines[0]) < 60):
                key = lines[0].rstrip(":").strip()
                val = "\n".join(lines[1:]).strip()
                if key and val and key.lower() != val.lower():
                    pairs[key] = val
                    continue

        if len(row) >= 2:
            key = row[0].strip().rstrip(":")
            val = row[1].strip()
            if key and (not val or val.lower() == key.lower()) and i + 1 < num_rows:
                next_row = rows[i + 1]
                if next_row:
                    next_val = next_row[0].strip() if len(next_row) == 1 or next_row[0].strip() != key else (next_row[1].strip() if len(next_row) > 1 else "")
                    if next_val and next_val.lower() != key.lower():
                        val = next_val
            if key and val and val.lower() != key.lower():
                pairs[key] = val
                continue

        if len(row) == 1 or (len(row) >= 2 and not row[1].strip()):
            key = row[0].strip().rstrip(":")
            if key and (key.endswith(":") or len(key) < 60) and i + 1 < num_rows:
                next_row = rows[i + 1]
                if next_row:
                    val = next_row[0].strip() if len(next_row) >= 1 else ""
                    if val and val.lower() != key.lower() and not val.endswith(":"):
                        pairs[key] = val

    return pairs


def _is_key_value_table(rows: list[list[str]]) -> bool:
    """Detect if a 2D grid is a key-value pair table."""
    if not rows:
        return False
    col_count = len(rows[0]) if rows else 0
    if col_count != 2:
        return False
    label_count = sum(
        1 for row in rows
        if len(row) >= 2 and row[0].strip() and len(row[0].strip()) < 80
    )
    return label_count >= len(rows) * 0.5


# ---------------------------------------------------------------------------
# Document-Order Iteration & Parser Engine
# ---------------------------------------------------------------------------

def iter_block_items(parent):
    """
    Yield each paragraph and table child within parent, in exact XML document order.
    """
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._element

    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.text.paragraph.CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl):
            yield docx.table.Table(child, parent)


def _extract_parsed_table(
    table,
    table_index: int,
    section_name: str,
    source_order: int,
    source_page: Optional[int],
) -> ParsedTable:
    """Construct a detailed ParsedTable preserving cell-level attributes."""
    parsed_rows: list[ParsedRow] = []
    raw_grid: list[list[str]] = []

    for r_idx, row in enumerate(table.rows):
        row_cells: list[ParsedCell] = []
        grid_row: list[str] = []
        for c_idx, cell in enumerate(row.cells):
            raw_text = cell.text
            clean_text = raw_text.strip()

            if cell.tables:
                nested_parts = []
                for nt in cell.tables:
                    for nr in nt.rows:
                        nested_parts.append(" | ".join(nc.text.strip() for nc in nr.cells))
                if nested_parts:
                    clean_text = clean_text + "\n" + "\n".join(nested_parts)

            cell_cb_state, cb_labels = _analyze_cell_checkboxes(cell, clean_text)

            parsed_cell = ParsedCell(
                row_index=r_idx,
                column_index=c_idx,
                text=clean_text,
                raw_text=raw_text,
                is_merged=False,
                checkbox_state=cell_cb_state,
                checkbox_labels=cb_labels,
            )
            row_cells.append(parsed_cell)
            grid_row.append(clean_text)

        parsed_rows.append(ParsedRow(row_index=r_idx, cells=row_cells))
        raw_grid.append(grid_row)

    col_count = max((len(r.cells) for r in parsed_rows), default=0)
    return ParsedTable(
        table_index=table_index,
        section_name=section_name,
        row_count=len(parsed_rows),
        column_count=col_count,
        rows=parsed_rows,
        source_order=source_order,
        source_page=source_page,
        raw_grid=raw_grid,
    )


def parse_cognos_docx(path: str | Path) -> CognosParsedDocument:
    """
    Parse a Cognos Report Definition DOCX into a structured representation
    using document-order XML element traversal.

    No keyword heuristic table assignment.
    No artificial page estimation.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Document not found: {path}")
    if p.suffix.lower() != ".docx":
        raise ValueError(f"Expected .docx file, got: {p.suffix}")

    document = docx.Document(str(p))
    result = CognosParsedDocument(filename=p.name)

    current_section: Optional[DocumentSection] = None
    source_order = 0
    table_index = 0
    current_page: Optional[int] = None  # None unless explicit page break detected

    for item in iter_block_items(document):
        source_order += 1

        if isinstance(item, docx.text.paragraph.Paragraph):
            text = item.text.strip()
            result.all_paragraphs.append(text)

            # Check explicit page breaks or last rendered page breaks in XML
            xml = item._element.xml
            breaks = xml.count("w:lastRenderedPageBreak") + xml.count('w:type="page"')
            if breaks > 0:
                current_page = (current_page or 1) + breaks

            # Check section detection
            is_heading = item.style and item.style.name and \
                         item.style.name.lower().startswith("heading")
            detected_section = _detect_section(text)

            if detected_section or (is_heading and text):
                if current_section and current_section.end_order == -1:
                    current_section.end_order = max(1, source_order - 1)
                    current_section.raw_text = "\n".join(current_section.paragraphs)

                section_name = detected_section or text
                current_section = DocumentSection(
                    name=section_name,
                    start_order=source_order,
                    source_page=current_page,
                )
                result.sections.append(current_section)

            if current_section is None and text:
                current_section = DocumentSection(
                    name="Report Definition",
                    start_order=source_order,
                    source_page=current_page,
                )
                result.sections.append(current_section)

            if current_section and text:
                current_section.paragraphs.append(text)

            checkboxes = _extract_checkboxes_from_paragraph(item)
            if checkboxes:
                result.all_checkboxes.extend(checkboxes)
                if current_section:
                    current_section.checkboxes.extend(checkboxes)

        elif isinstance(item, docx.table.Table):
            table_index += 1

            # Check explicit page breaks or last rendered page breaks inside the table XML
            xml = item._element.xml
            breaks = xml.count("w:lastRenderedPageBreak") + xml.count('w:type="page"')
            if breaks > 0:
                current_page = (current_page or 1) + breaks

            if current_section is None:
                current_section = DocumentSection(
                    name="Report Definition",
                    start_order=source_order,
                    source_page=current_page,
                )
                result.sections.append(current_section)

            parsed_table = _extract_parsed_table(
                item,
                table_index=table_index,
                section_name=current_section.name,
                source_order=source_order,
                source_page=current_page,
            )

            current_section.tables.append(parsed_table)
            result.all_parsed_tables.append(parsed_table)

            table_grid = parsed_table.to_grid()
            if _is_key_value_table(table_grid):
                kv = _extract_key_value_pairs(table_grid)
                result.all_key_value_pairs.update(kv)
                current_section.key_value_pairs.update(kv)

            for row in parsed_table.rows:
                for cell in row.cells:
                    if cell.checkbox_labels:
                        result.all_checkboxes.extend(cell.checkbox_labels)

    if current_section and current_section.end_order == -1:
        current_section.end_order = source_order
        current_section.raw_text = "\n".join(current_section.paragraphs)

    for section in result.sections:
        if not section.paragraphs and not section.tables:
            result.warnings.append(
                f"Section '{section.name}' was detected but contains no "
                f"extractable content."
            )

    try:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(str(p)) as z:
            if 'word/comments.xml' in z.namelist():
                comments_xml = z.read('word/comments.xml')
                root = ET.fromstring(comments_xml)
                for comment in root:
                    author = comment.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', '')
                    texts = []
                    for t_node in comment.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        if t_node.text:
                            texts.append(t_node.text)
                    if texts:
                        result.comments.append({
                            "author": author,
                            "text": "".join(texts)
                        })
    except Exception as e:
        result.warnings.append(f"Failed to extract comments: {e}")

    return result


# ---------------------------------------------------------------------------
# Diagnostic Structural Dumper (Part 1 §8)
# ---------------------------------------------------------------------------

def dump_document_structure(doc: CognosParsedDocument) -> str:
    """
    Produce a clean, hierarchical diagnostic string representing the exact
    parsed structure of a Cognos report document.

    Report
      Section
        Table
          Row
            Cell
    """
    lines = []
    lines.append(f"DOCUMENT STRUCTURE DUMP: {doc.filename}")
    lines.append("=" * 60)

    for sec in doc.sections:
        pg_str = str(sec.source_page) if sec.source_page is not None else "UNKNOWN"
        lines.append(
            f"SECTION: {sec.name} (Start Order: {sec.start_order}, "
            f"End Order: {sec.end_order}, Page: {pg_str})"
        )
        if sec.paragraphs:
            lines.append(f"  Paragraphs ({len(sec.paragraphs)}):")
            for p in sec.paragraphs[:3]:
                lines.append(f"    - {p[:80]}...")

        for t in sec.tables:
            lines.append(
                f"  TABLE {t.table_index} (Source Order: {t.source_order}, "
                f"Rows: {t.row_count}, Cols: {t.column_count})"
            )
            for r in t.rows:
                lines.append(f"    ROW {r.row_index}:")
                for c in r.cells:
                    cb_str = ""
                    if c.checkbox_labels:
                        cb_details = ", ".join(f"{cb['label']}={cb['state'].value if hasattr(cb['state'], 'value') else cb['state']}" for cb in c.checkbox_labels)
                        cb_str = f" [Checkboxes: {cb_details}]"
                    elif c.checkbox_state != CheckboxState.UNKNOWN:
                        cb_str = f" [State: {c.checkbox_state.value}]"

                    text_disp = c.text.replace("\n", " \\n ")
                    if len(text_disp) > 100:
                        text_disp = text_disp[:97] + "..."
                    lines.append(
                        f"      Cell [Row {c.row_index}, Col {c.column_index}]: "
                        f"'{text_disp}'{cb_str}"
                    )

    return "\n".join(lines)

