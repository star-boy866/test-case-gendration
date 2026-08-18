import docx
import sys
import json
import os

def dump_docx(file_path):
    doc = docx.Document(file_path)
    output = {
        'paragraphs': [],
        'tables': []
    }
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            output['paragraphs'].append({
                'index': i,
                'text': text
            })
            
    for i, table in enumerate(doc.tables):
        t_data = []
        for r_idx, row in enumerate(table.rows):
            r_data = []
            for c_idx, cell in enumerate(row.cells):
                r_data.append(cell.text.strip())
            t_data.append(r_data)
        output['tables'].append({
            'index': i,
            'rows': t_data
        })
        
    return output

if __name__ == '__main__':
    if len(sys.argv) > 1:
        res = dump_docx(sys.argv[1])
        with open('poc/document_parsing_benchmark/scratch_dump.json', 'w', encoding='utf-8') as f:
            json.dump(res, f, indent=2, ensure_ascii=False)
        print("Dumped to poc/document_parsing_benchmark/scratch_dump.json")
