"""
Document parser — Phase 1.

Turns uploaded RDD/LDM files (.xlsx, .csv, .docx, .pdf) into structured
records ready for Knowledge Base persistence.

HALLUCINATION PREVENTION IS THE CORE DESIGN CONSTRAINT HERE:
- We only extract what is EXPLICITLY present in a recognizable tabular
  structure (a sheet, a docx table, or a PDF table) with headers that map
  to known canonical fields (table name, column name, data type, key type,
  join mapping, valid value, business rule).
- Free-form prose (docx paragraphs, PDF body text with no table structure)
  is NEVER auto-classified as a business rule or any other KB fact. It is
  captured as an `UnstructuredNote` for human review only.
- If nothing recognizable is found anywhere in the document, the caller
  (ingestion API) must surface the mandated message:
  "Insufficient metadata available. Additional documentation is required
  before generation can continue."
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import pandas as pd

SUPPORTED_EXTENSIONS = {".xlsx", ".xls", ".csv", ".docx", ".pdf"}

# --- Canonical field synonym map -------------------------------------------------
# Real-world LDMs/RDDs use inconsistent header naming. We normalize aggressively
# but ONLY match against this explicit allow-list — no fuzzy free-for-all that
# could misclassify unrelated columns.

_SYNONYMS: dict[str, list[str]] = {
    "table_name": ["table name", "table", "entity", "entity name", "table_name"],
    "column_name": ["column name", "column", "field", "field name", "attribute", "column_name"],
    "data_type": ["data type", "datatype", "type", "data_type"],
    "key_type": ["key", "key type", "pk/fk", "pk_fk", "primary/foreign key", "key_type"],
    "description": ["description", "definition", "notes", "comments"],
    "from_table": ["from table", "source table", "parent table", "from_table"],
    "from_column": ["from column", "source column", "parent column", "from_column"],
    "to_table": ["to table", "target table", "child table", "to_table"],
    "to_column": ["to column", "target column", "child column", "to_column"],
    "join_type": ["join type", "join_type"],
    "valid_value": ["valid value", "valid values", "allowed value", "allowed values", "value", "valid_value"],
    "meaning": ["meaning", "value description", "value meaning"],
    "rule_text": ["business rule", "rule", "rule description", "rule_text", "validation rule"],
}


def _normalize_header(h: str) -> str:
    return re.sub(r"\s+", " ", str(h).strip().lower())


def _map_headers(columns: list[str]) -> dict[str, str]:
    """Return {actual_column_name: canonical_field_name} for recognized headers only."""
    mapping: dict[str, str] = {}
    normalized = {col: _normalize_header(col) for col in columns}
    for canonical, synonyms in _SYNONYMS.items():
        for col, norm in normalized.items():
            if norm in synonyms and col not in mapping:
                mapping[col] = canonical
    return mapping


@dataclass
class ParsedDocument:
    tables: list[dict] = field(default_factory=list)         # table_name, description
    columns: list[dict] = field(default_factory=list)        # table_name, column_name, data_type, key_type, description
    joins: list[dict] = field(default_factory=list)          # from_table, from_column, to_table, to_column, join_type
    valid_values: list[dict] = field(default_factory=list)   # table_name, column_name, valid_value, meaning
    business_rules: list[dict] = field(default_factory=list) # rule_text, related_table, related_column
    unstructured_notes: list[dict] = field(default_factory=list)  # content, reason
    warnings: list[str] = field(default_factory=list)
    sheets_parsed: int = 0
    sheets_skipped: int = 0

    @property
    def has_structured_content(self) -> bool:
        return bool(
            self.tables or self.columns or self.joins
            or self.valid_values or self.business_rules
        )


def _classify_and_extract(df: pd.DataFrame, sheet_label: str, result: ParsedDocument) -> bool:
    """
    Given a dataframe with a header row, figure out which canonical schema it
    matches (column-level LDM, join mapping, valid-values, or business rules)
    and extract accordingly. Returns True if anything was extracted.
    """
    if df.empty or len(df.columns) == 0:
        return False

    header_map = _map_headers([str(c) for c in df.columns])
    canonical_fields = set(header_map.values())
    if not canonical_fields:
        return False

    # Rename columns to canonical names for easy access; keep only mapped ones.
    df = df.rename(columns=header_map)
    df = df[[c for c in df.columns if c in _SYNONYMS.keys()]]
    df = df.dropna(how="all")

    extracted_any = False

    # Classification is MUTUALLY EXCLUSIVE per sheet, most-specific signature
    # first. This prevents a valid-values or business-rules sheet that also
    # happens to carry table_name/column_name context columns from being
    # double-counted as a column-level LDM sheet too.

    is_join_sheet = {"from_table", "from_column", "to_table", "to_column"}.issubset(canonical_fields)
    is_valid_value_sheet = (not is_join_sheet) and "valid_value" in canonical_fields
    is_rule_sheet = (not is_join_sheet) and (not is_valid_value_sheet) and "rule_text" in canonical_fields
    is_column_ldm_sheet = (
        not is_join_sheet and not is_valid_value_sheet and not is_rule_sheet
        and {"table_name", "column_name"}.issubset(canonical_fields)
    )

    if is_join_sheet:
        for _, row in df.iterrows():
            if pd.isna(row.get("from_table")) or pd.isna(row.get("to_table")):
                continue
            result.joins.append({
                "from_table": str(row["from_table"]).strip(),
                "from_column": str(row.get("from_column", "")).strip(),
                "to_table": str(row["to_table"]).strip(),
                "to_column": str(row.get("to_column", "")).strip(),
                "join_type": (str(row["join_type"]).strip() if "join_type" in df.columns and pd.notna(row.get("join_type")) else None),
            })
            extracted_any = True

    elif is_valid_value_sheet:
        for _, row in df.iterrows():
            if pd.isna(row.get("valid_value")):
                continue
            result.valid_values.append({
                "table_name": (str(row["table_name"]).strip() if "table_name" in df.columns and pd.notna(row.get("table_name")) else ""),
                "column_name": (str(row["column_name"]).strip() if "column_name" in df.columns and pd.notna(row.get("column_name")) else ""),
                "valid_value": str(row["valid_value"]).strip(),
                "meaning": (str(row["meaning"]).strip() if "meaning" in df.columns and pd.notna(row.get("meaning")) else None),
            })
            extracted_any = True

    elif is_rule_sheet:
        for _, row in df.iterrows():
            if pd.isna(row.get("rule_text")) or not str(row["rule_text"]).strip():
                continue
            result.business_rules.append({
                "rule_text": str(row["rule_text"]).strip(),
                "related_table": (str(row["table_name"]).strip() if "table_name" in df.columns and pd.notna(row.get("table_name")) else None),
                "related_column": (str(row["column_name"]).strip() if "column_name" in df.columns and pd.notna(row.get("column_name")) else None),
            })
            extracted_any = True

    elif is_column_ldm_sheet:
        seen_tables: dict[str, Optional[str]] = {}
        for _, row in df.iterrows():
            if pd.isna(row.get("table_name")) or pd.isna(row.get("column_name")):
                continue
            t_name = str(row["table_name"]).strip()
            c_name = str(row["column_name"]).strip()
            desc = str(row["description"]).strip() if "description" in df.columns and pd.notna(row.get("description")) else None

            if t_name not in seen_tables:
                seen_tables[t_name] = desc
            result.columns.append({
                "table_name": t_name,
                "column_name": c_name,
                "data_type": (str(row["data_type"]).strip() if "data_type" in df.columns and pd.notna(row.get("data_type")) else None),
                "key_type": (str(row["key_type"]).strip() if "key_type" in df.columns and pd.notna(row.get("key_type")) else None),
                "description": desc,
            })
            extracted_any = True

        for t_name, desc in seen_tables.items():
            result.tables.append({"table_name": t_name, "description": desc})

    return extracted_any


def _parse_dataframe_into(df: pd.DataFrame, sheet_label: str, result: ParsedDocument) -> None:
    if df.empty:
        result.sheets_skipped += 1
        return
    extracted = _classify_and_extract(df, sheet_label, result)
    if extracted:
        result.sheets_parsed += 1
    else:
        result.sheets_skipped += 1
        result.warnings.append(
            f"Sheet/table '{sheet_label}' did not match any recognized LDM/RDD "
            f"column pattern and was skipped (no data invented)."
        )


def parse_xlsx(path: str | Path) -> ParsedDocument:
    result = ParsedDocument()
    xls = pd.ExcelFile(path)
    for sheet_name in xls.sheet_names:
        df = xls.parse(sheet_name, dtype=str)
        _parse_dataframe_into(df, sheet_name, result)
    return result


def parse_csv(path: str | Path) -> ParsedDocument:
    result = ParsedDocument()
    df = pd.read_csv(path, dtype=str)
    _parse_dataframe_into(df, Path(path).name, result)
    return result


def parse_docx(path: str | Path) -> ParsedDocument:
    import docx  # python-docx

    result = ParsedDocument()
    document = docx.Document(str(path))

    for idx, table in enumerate(document.tables):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if len(rows) < 2:
            result.sheets_skipped += 1
            continue
        header, *body = rows
        df = pd.DataFrame(body, columns=header)
        _parse_dataframe_into(df, f"table_{idx + 1}", result)

    # Paragraph text is captured for human review only — never auto-parsed
    # into business rules or any structured KB fact.
    paragraph_text = "\n".join(p.text.strip() for p in document.paragraphs if p.text.strip())
    if paragraph_text:
        result.unstructured_notes.append({
            "content": paragraph_text[:5000],  # cap to keep DB rows sane
            "reason": "Free-form docx paragraph text — not auto-structured; requires human review.",
        })

    return result


def parse_pdf(path: str | Path) -> ParsedDocument:
    import pdfplumber

    result = ParsedDocument()
    unstructured_chunks: list[str] = []

    with pdfplumber.open(str(path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            tables = page.extract_tables()
            if tables:
                for t_idx, table in enumerate(tables):
                    if not table or len(table) < 2:
                        continue
                    header, *body = table
                    df = pd.DataFrame(body, columns=header)
                    _parse_dataframe_into(df, f"page{page_num}_table{t_idx + 1}", result)
            else:
                text = page.extract_text()
                if text and text.strip():
                    unstructured_chunks.append(f"[Page {page_num}]\n{text.strip()}")

    if unstructured_chunks:
        result.unstructured_notes.append({
            "content": "\n\n".join(unstructured_chunks)[:5000],
            "reason": "PDF page text without recognizable table structure — requires human review.",
        })

    return result


def parse_document(path: str | Path) -> ParsedDocument:
    """Dispatch by extension. Raises ValueError for unsupported types."""
    p = Path(path)
    ext = p.suffix.lower()

    if ext in (".xlsx", ".xls"):
        return parse_xlsx(p)
    if ext == ".csv":
        return parse_csv(p)
    if ext == ".docx":
        return parse_docx(p)
    if ext == ".pdf":
        return parse_pdf(p)

    raise ValueError(
        f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )
