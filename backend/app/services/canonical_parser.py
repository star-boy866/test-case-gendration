import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from dataclasses import dataclass, field
from typing import Dict, List, Any, Optional
import docx

namespaces = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'w15': 'http://schemas.microsoft.com/office/word/2012/wordml'
}

@dataclass
class CanonicalTableCell:
    text: str
    row_span: int = 1
    col_span: int = 1

@dataclass
class CanonicalTableRow:
    cells: List[CanonicalTableCell]

@dataclass
class CanonicalTable:
    index: int
    rows: List[CanonicalTableRow]

@dataclass
class CanonicalDocumentModel:
    tables: List[CanonicalTable] = field(default_factory=list)
    paragraphs: List[str] = field(default_factory=list)
    checkboxes: List[Dict[str, Any]] = field(default_factory=list)
    comments: List[Dict[str, Any]] = field(default_factory=list)

def _extract_ooxml_supplements(docx_path: str) -> tuple[List[Dict], List[Dict]]:
    comments = []
    checkboxes = []
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            if 'word/comments.xml' in docx_zip.namelist():
                comments_xml = docx_zip.read('word/comments.xml')
                root = ET.fromstring(comments_xml)
                for comment in root.findall('.//w:comment', namespaces):
                    c_id = comment.get(f'{{{namespaces["w"]}}}id')
                    author = comment.get(f'{{{namespaces["w"]}}}author', 'Unknown')
                    texts = [t.text for t in comment.findall('.//w:t', namespaces) if t.text]
                    comments.append({
                        'id': c_id,
                        'author': author,
                        'text': ''.join(texts)
                    })
            
            if 'word/document.xml' in docx_zip.namelist():
                doc_xml = docx_zip.read('word/document.xml')
                root = ET.fromstring(doc_xml)
                
                # Form field checkboxes
                for ffdata in root.findall('.//w:ffData', namespaces):
                    name_el = ffdata.find('w:name', namespaces)
                    cb_el = ffdata.find('w:checkBox', namespaces)
                    if cb_el is not None:
                        checked_el = cb_el.find('w:checked', namespaces)
                        default_el = cb_el.find('w:default', namespaces)
                        is_checked = False
                        if checked_el is not None:
                            val = checked_el.get(f'{{{namespaces["w"]}}}val')
                            is_checked = (val in ['1', 'true', 'True'])
                        elif default_el is not None:
                            val = default_el.get(f'{{{namespaces["w"]}}}val')
                            is_checked = (val in ['1', 'true', 'True'])
                        
                        name = name_el.get(f'{{{namespaces["w"]}}}val') if name_el is not None else 'unknown'
                        checkboxes.append({'type': 'legacy', 'name': name, 'checked': is_checked})
                
                # Content control checkboxes
                for sdt in root.findall('.//w:sdt', namespaces):
                    cb = sdt.find('.//w14:checkbox', namespaces)
                    if cb is not None:
                        checked_el = cb.find('w14:checked', namespaces)
                        is_checked = False
                        if checked_el is not None:
                            val = checked_el.get(f'{{{namespaces["w14"]}}}val')
                            is_checked = (val in ['1', 'true', 'True'])
                        alias_el = sdt.find('.//w:alias', namespaces)
                        name = alias_el.get(f'{{{namespaces["w"]}}}val') if alias_el is not None else 'content_control'
                        checkboxes.append({'type': 'content_control', 'name': name, 'checked': is_checked})
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning(f"Error extracting OOXML supplements: {e}")
        
    return comments, checkboxes

def parse_canonical_docx(path: str | Path) -> CanonicalDocumentModel:
    path_str = str(path)
    model = CanonicalDocumentModel()
    
    # 1. Structural extraction via python-docx
    doc = docx.Document(path_str)
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            model.paragraphs.append(text)
            
    for i, table in enumerate(doc.tables):
        canonical_rows = []
        for row in table.rows:
            canonical_cells = []
            for cell in row.cells:
                # Basic representation. python-docx doesn't easily expose rowspan natively 
                # without iterating cell._tc, but we store the string contents cleanly.
                text = cell.text.strip()
                canonical_cells.append(CanonicalTableCell(text=text))
            canonical_rows.append(CanonicalTableRow(cells=canonical_cells))
        model.tables.append(CanonicalTable(index=i, rows=canonical_rows))
        
    # 2. OOXML Supplement
    comments, checkboxes = _extract_ooxml_supplements(path_str)
    model.comments = comments
    model.checkboxes = checkboxes
    
    return model
